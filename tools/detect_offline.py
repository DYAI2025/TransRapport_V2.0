#!/usr/bin/env python3
"""
Offline detect/report generator from data/sessions/<sid>/segments.json

Usage:
  python tools/detect_offline.py --sid <session_id>
"""
import argparse
import json
from pathlib import Path
from typing import Dict, List, Any

import sys
from pathlib import Path as _Path
# ensure repo root on path for 'engine' and 'reports'
sys.path.insert(0, str(_Path(__file__).resolve().parents[1]))

from engine.topics import extract_topics
from engine.chapters import detect_chapters
from engine.summarizer import Summarizer
from reports.renderer import render_markdown


def _compute_stats(segments: List[Dict[str, Any]]) -> Dict[str, Any]:
    totals: Dict[str, float] = {}
    total_dur = 0.0
    for s in segments:
        dur = max(float(s.get('t1', 0.0)) - float(s.get('t0', 0.0)), 0.0)
        spk = s.get('speaker') or 'UNK'
        totals[spk] = totals.get(spk, 0.0) + dur
        total_dur += dur
    shares = {k: (v / total_dur if total_dur > 0 else 0.0) for k, v in totals.items()}
    return {
        'speakers': [{'speaker': k, 'seconds': round(v, 3), 'share': round(shares[k], 3)} for k, v in sorted(totals.items())],
        'total_seconds': round(total_dur, 3),
    }


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--sid", required=True)
    args = ap.parse_args()

    sid = args.sid
    sess_dir = Path("data/sessions") / sid
    seg_path = sess_dir / "segments.json"
    if not seg_path.exists():
        raise SystemExit(f"No segments.json at {seg_path}")
    segments: List[Dict[str, Any]] = json.loads(seg_path.read_text(encoding="utf-8"))

    # Topics/Chapters/Summaries
    topics = extract_topics(segments, max_topics=8)
    chapters = detect_chapters(segments)
    summ = Summarizer(mode=None)
    global_summary = summ.summarize(segments, lang=None)
    per_ch_summ: List[str] = []
    for ch in chapters:
        i0, i1 = int(ch.get('start_idx', 0)), int(ch.get('end_idx', 0))
        per_ch_summ.append(summ.summarize(segments[i0:i1], lang=None))

    md = render_markdown(sid, segments, topics, chapters, global_summary, per_ch_summ)
    # Save MD
    (sess_dir / 'report.md').write_text(md, encoding='utf-8')
    # Save TXT (simple)
    txt_lines = []
    for line in md.splitlines():
        l = line.lstrip('#').lstrip().replace("- ", "• ")
        txt_lines.append(l)
    (sess_dir / 'report.txt').write_text("\n".join(txt_lines) + "\n", encoding='utf-8')

    # DOCX
    try:
        from docx import Document  # type: ignore
        doc = Document()
        for line in md.splitlines():
            if line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('- '):
                p = doc.add_paragraph()
                p.add_run('• ').bold = False
                p.add_run(line[2:])
            else:
                doc.add_paragraph(line)
        doc.save(str(sess_dir / 'report.docx'))
    except Exception:
        pass

    # PDF
    try:
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.pdfgen import canvas  # type: ignore
        from reportlab.lib.units import cm  # type: ignore
        c = canvas.Canvas(str(sess_dir / 'report.pdf'), pagesize=A4)
        width, height = A4
        x, y = 2 * cm, height - 2 * cm
        for raw in md.splitlines():
            line = raw.replace('\t','    ')
            if not line:
                y -= 0.5 * cm
            else:
                c.drawString(x, y, line[:120])
                y -= 0.6 * cm
            if y < 2 * cm:
                c.showPage(); y = height - 2 * cm
        c.save()
    except Exception:
        pass

    # Save detect report
    out = {
        "session_id": sid,
        "segments": len(segments),
        "stats": _compute_stats(segments),
        "topics": topics,
        "chapters": chapters,
        "report_files": {
            "md": str(sess_dir / 'report.md'),
            "txt": str(sess_dir / 'report.txt'),
            "docx": str(sess_dir / 'report.docx'),
            "pdf": str(sess_dir / 'report.pdf'),
        },
    }
    (sess_dir / 'detect_offline.json').write_text(json.dumps(out, indent=2, ensure_ascii=False), encoding='utf-8')
    print(json.dumps({"ok": True, "out": str(sess_dir / 'detect_offline.json')}, ensure_ascii=False))


if __name__ == '__main__':
    main()
