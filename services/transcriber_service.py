import os
import json
import uuid
import queue
import threading
import tempfile
import time
from datetime import datetime
import re
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple

import yaml
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, BackgroundTasks
from fastapi.responses import JSONResponse, PlainTextResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from pydantic import ConfigDict, PrivateAttr, Field
try:
    from faster_whisper import WhisperModel  # type: ignore
except Exception:  # pragma: no cover - allow running in fake mode without the package
    WhisperModel = None  # type: ignore

try:
    from .stt_whisper_ct2 import WhisperCT2Pipeline
    _whisper_pipeline = None
except ImportError:
    WhisperCT2Pipeline = None
    _whisper_pipeline = None

from .srt_export import segments_to_srt, segments_to_txt
from engine.diarization import DiarizerFactory
from engine.marker_engine_core import MarkerBundle, MarkerEngine
from engine.scoring_engine import IntuitionTelemetry
from engine.prosody import extract_prosody
from engine.summarizer import Summarizer
from engine.topics import extract_topics
from engine.chapters import detect_chapters
from engine.logging_framework import get_logger, with_error_handling, setup_common_error_handlers
from engine.report_pipeline import ReportGenerator, ReportRequest, ReportFormat, generate_report_background

# Disable HF Xet optimized downloads to avoid native dependency issues
os.environ.setdefault("HF_HUB_ENABLE_HF_XET", "0")


# ---- Config loading ----
def _load_config() -> Dict[str, Any]:
    cfg_path = Path("config/app.yaml")
    if cfg_path.exists():
        try:
            with cfg_path.open("r", encoding="utf-8") as f:
                return yaml.safe_load(f) or {}
        except Exception:
            return {}
    return {}


_CFG = _load_config()

# Prefer env vars; fall back to config file; then to sane defaults
DATA_ROOT = Path(os.environ.get(
    "SERAPI_DATA_ROOT",
    _CFG.get("paths", {}).get("data_root", "./data/sessions")
))
MODEL_DIR = os.environ.get(
    "SERAPI_WHISPER_MODEL_DIR",
    _CFG.get("paths", {}).get("whisper_model_dir", "./models/whisper/faster-whisper-large-v3")
)
def _norm_lang(v: Optional[str]) -> Optional[str]:
    if v is None:
        return None
    s = str(v).strip().lower()
    if s in {"", "auto", "none", "null"}:
        return None
    return s

LANG_DEFAULT = _norm_lang(os.environ.get(
    "SERAPI_LANG_DEFAULT",
    _CFG.get("runtime", {}).get("language_default")
))
WINDOW_SECONDS = int(_CFG.get("runtime", {}).get("window_seconds", 25))
IDLE_FLUSH_SECONDS = float(_CFG.get("runtime", {}).get("idle_flush_seconds", 2.0))
FAKE_ASR = os.environ.get("SERAPI_FAKE_ASR", "").lower() in {"1", "true", "yes"}
MODEL_ID = os.environ.get("SERAPI_MODEL_ID", "large-v3")
ALLOW_DOWNLOAD = os.environ.get("SERAPI_ALLOW_DOWNLOAD", "").lower() in {"1", "true", "yes"}
DEMO_MARKERS = os.environ.get("SERAPI_DEMO_MARKERS", "").lower() in {"1", "true", "yes"}


app = FastAPI(title="Serapi Transcriber (I1) - Offline")

# Initialize logging and error handling
setup_common_error_handlers()
logger = get_logger('transcriber_service')

# Initialize report generator
report_generator = ReportGenerator(DATA_ROOT)


# ---- In-Memory Session Registry ----
class Session(BaseModel):
    model_config = ConfigDict(arbitrary_types_allowed=True)

    id: str
    lang: Optional[str] = None
    segments: List[Dict[str, Any]] = Field(default_factory=list)

    # Internal (excluded from JSON)
    _pcm_queue: queue.Queue = PrivateAttr(default_factory=queue.Queue)
    _stop_flag: bool = PrivateAttr(default=False)
    _worker: Optional[threading.Thread] = PrivateAttr(default=None)
    _subscribers: List[Tuple[Any, Any]] = PrivateAttr(default_factory=list)  # (loop, asyncio.Queue)
    _diarizer: object = PrivateAttr(default=None)
    _marker_engine: Optional[MarkerEngine] = PrivateAttr(default=None)
    _telemetry: Optional[IntuitionTelemetry] = PrivateAttr(default=None)

    def add_subscriber(self, loop, async_queue) -> None:
        self._subscribers.append((loop, async_queue))

    def remove_subscriber(self, async_queue) -> None:
        self._subscribers = [(l, q) for (l, q) in self._subscribers if q is not async_queue]


SESSIONS: Dict[str, Session] = {}


# ---- Ensure directories ----
def ensure_session_dir(session_id: str) -> Path:
    d = DATA_ROOT / session_id
    d.mkdir(parents=True, exist_ok=True)
    return d


def save_segments(session: Session):
    d = ensure_session_dir(session.id)
    with (d / "segments.json").open("w", encoding="utf-8") as f:
        json.dump(session.segments, f, ensure_ascii=False, indent=2)


# ---- Load Whisper Model once (or run in fake mode) ----
whisper_model = None
RUN_MODE = "mock"
MODEL_ERROR: Optional[str] = None
if not FAKE_ASR and WhisperModel is not None:
    try:
        model_source = None
        if os.path.isdir(MODEL_DIR):
            model_source = MODEL_DIR
        elif ALLOW_DOWNLOAD:
            model_source = MODEL_ID  # triggers auto-download to HF cache
        if model_source is not None:
            try:
                whisper_model = WhisperModel(model_source, device="auto", compute_type="int8_float16")
            except Exception:
                whisper_model = WhisperModel(model_source, device="auto", compute_type="int8")
            RUN_MODE = "real"
    except Exception as e:
        MODEL_ERROR = str(e)
        whisper_model = None
        RUN_MODE = "mock"


def _fake_transcribe_window(pcm_bytes: bytes, sr: int):
    # Estimate seconds from PCM16 mono byte length
    n_bytes = len(pcm_bytes)
    sec = max(n_bytes / float(2 * sr), 0.0)
    # Emit a simple single segment covering the chunk
    return [{
        "t0": 0.0,
        "t1": round(sec, 3),
        "text": "[mock transcript]",
        "words": [],
    }]


def _init_whisper_pipeline():
    """Initialize the Whisper CT2 pipeline if available."""
    global _whisper_pipeline
    if WhisperCT2Pipeline is not None and _whisper_pipeline is None:
        model_dir = None
        if MODEL_DIR and hasattr(MODEL_DIR, 'exists') and MODEL_DIR.exists():
            model_dir = str(MODEL_DIR)
        _whisper_pipeline = WhisperCT2Pipeline(
            model_dir=model_dir,
            mock_mode=FAKE_ASR
        )
    return _whisper_pipeline


def transcribe_window(pcm_bytes: bytes, sr: int, lang: Optional[str]):
    """
    Transkribiert einen PCM16-Mono-Chunk (bytes) mit faster-whisper oder CT2 pipeline.
    """
    # Try using the new CT2 pipeline first
    pipeline = _init_whisper_pipeline()
    if pipeline is not None:
        try:
            import numpy as np
            # Convert PCM16 bytes to float32 array
            pcm_array = np.frombuffer(pcm_bytes, dtype="<i2").astype("float32") / 32768.0
            segments = pipeline.transcribe_array(pcm_array, sr, language=lang)
            return segments
        except Exception as e:
            logger.warning(f"CT2 pipeline failed, falling back to legacy: {e}")
    
    # Fallback to legacy implementation
    if FAKE_ASR or whisper_model is None:
        return _fake_transcribe_window(pcm_bytes, sr)

    import numpy as np
    import soundfile as sf

    # Bytes -> np.int16 -> float32 / -1..1
    pcm = np.frombuffer(pcm_bytes, dtype="<i2").astype("float32") / 32768.0

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tmp:
            tmp_path = tmp.name
            sf.write(tmp.name, pcm, sr)
        # Normalize language (None = auto)
        language = _norm_lang(lang)
        try:
            segments, info = whisper_model.transcribe(
                tmp_path,
                language=language,
                vad_filter=True,
                word_timestamps=True,
            )
        except Exception:
            # If invalid language slipped through, retry with auto
            segments, info = whisper_model.transcribe(
                tmp_path,
                language=None,
                vad_filter=True,
                word_timestamps=True,
            )
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

    out = []
    for s in segments:
        out.append({
            "t0": float(getattr(s, "start", 0.0)),
            "t1": float(getattr(s, "end", 0.0)),
            "text": (getattr(s, "text", "") or "").strip(),
            "words": [
                {
                    "w": getattr(w, "word", ""),
                    "t0": float(getattr(w, "start", getattr(s, "start", 0.0)) or getattr(s, "start", 0.0)),
                    "t1": float(getattr(w, "end", getattr(s, "end", 0.0)) or getattr(s, "end", 0.0)),
                }
                for w in (getattr(s, "words", None) or [])
            ],
        })
    return out


FLUSH_SENTINEL = object()


def _worker_loop(session: Session, sr: int, window_sec: int, loop, out_async_queue, idle_flush_sec: float = 2.0, soft_flush_sec: float = 0.0):
    """
    Worker-Thread sammelt PCM, transkribiert bei Fenstergröße und postet
    Partial-Events in eine asyncio.Queue des WS-Event-Loops.
    """
    buf = bytearray()
    bytes_per_sec = 2 * sr  # 16-bit mono
    target_bytes = int(window_sec * bytes_per_sec)
    soft_target = int(max(0.0, soft_flush_sec) * bytes_per_sec)
    last_rx = time.time()

    def _broadcast(payload: dict):
        try:
            import asyncio
            # to local WS channel
            asyncio.run_coroutine_threadsafe(out_async_queue.put(payload), loop)
            # to any subscribers
            for sub_loop, sub_q in list(session._subscribers):
                try:
                    asyncio.run_coroutine_threadsafe(sub_q.put(payload), sub_loop)
                except Exception:
                    pass
        except Exception:
            pass

    def detect_demo_markers(text: str) -> List[str]:
        if not DEMO_MARKERS:
            return []
        t = (text or "").lower()
        hits: List[str] = []
        # guilt framing (avoid 'leider' false positives)
        if re.search(r"\b(es\s+tut\s+mir\s+leid|tut\s+mir\s+leid|entschuldigung|sorry|schuld)\b", t):
            hits.append("SEM_GUILT_FRAMING")
        if any(k in t for k in ["unsicher", "weiß nicht", "weiss nicht", "vielleicht", "unsure", "maybe"]):
            hits.append("SEM_DOUBT_UNCERTAINTY")
        if any(k in t for k in ["versprechen", "commit", "zusage", "verbind"]):
            hits.append("SEM_COMMITMENT_REQUEST")
        if any(k in t for k in ["streit", "konflikt", "ärger", "wuetend", "wütend", "angry"]):
            hits.append("CLU_INTUITION_CONFLICT")
        return hits

    def do_flush():
        nonlocal buf
        if not buf:
            return
        try:
            pcm_bytes = bytes(buf)
            new_segments = transcribe_window(pcm_bytes, sr, session.lang or LANG_DEFAULT)
            # Diarization: assign speaker labels per new segment
            try:
                import numpy as np
                pcm_f32 = np.frombuffer(pcm_bytes, dtype='<i2').astype('float32') / 32768.0
                labels = session._diarizer.label_segments(pcm_f32, sr, new_segments)
                # Prosody per segment (if not already provided by CT2 pipeline)
                for seg in new_segments:
                    if 'prosody' not in seg or not seg['prosody']:
                        try:
                            seg['prosody'] = extract_prosody(pcm_f32, sr, float(seg.get('t0', 0.0)), float(seg.get('t1', 0.0)))
                        except Exception:
                            seg['prosody'] = {}
                for seg, lab in zip(new_segments, labels):
                    seg['speaker'] = lab
            except Exception:
                pass
            session.segments.extend(new_segments)
            save_segments(session)

            payload = {
                "type": "partial_transcript",
                "added": new_segments,
                "total_segments": len(session.segments),
            }
            _broadcast(payload)

            # Demo marker events (optional)
            if DEMO_MARKERS:
                for seg in new_segments:
                    marks = detect_demo_markers(seg.get("text", ""))
                    for name in marks:
                        fam = "CLU" if name.startswith("CLU_") else ("SEM" if name.startswith("SEM_") else ("ATO" if name.startswith("ATO_") else "MARK"))
                        _broadcast({
                            "type": "marker",
                            "family": fam,
                            "name": name,
                            "t0": seg.get("t0", 0.0),
                            "t1": seg.get("t1", 0.0),
                            "text": seg.get("text", ""),
                        })
                        # feed telemetry for demo CLU_INTUITION as well
                        if session._telemetry and str(name).startswith("CLU_INTUITION"):
                            try:
                                # use current message index from marker engine if present, else fallback to 0
                                msg_idx = session._marker_engine.msg_idx if getattr(session, "_marker_engine", None) else 0
                                session._telemetry.on_provisional(str(name), int(msg_idx))
                            except Exception:
                                pass
            # Real markers (I3 MVP): example-based + ATO heuristics + CLU windows
            if session._marker_engine:
                try:
                    events = session._marker_engine.process(new_segments, session.segments)
                    for ev in events:
                        name = str(ev.get("name", ""))
                        fam = str(ev.get("family", ""))
                        if session._telemetry and name.startswith("CLU_INTUITION"):
                            # First update telemetry, then emit a single event (confirmed if window satisfied)
                            msg_idx = int(ev.get("msg_idx", 0))
                            session._telemetry.on_provisional(name, msg_idx)
                            will_confirm = False
                            try:
                                st = session._telemetry.states.get(name) or {}
                                will_confirm = (st.get("last_confirmed_at_msg_idx") == msg_idx) and bool(st.get("multiplier_active"))
                            except Exception:
                                will_confirm = False
                            _broadcast({
                                "type": "marker",
                                "family": fam or "CLU",
                                "name": name,
                                "confirmed": will_confirm or bool(ev.get("confirmed")),
                                "t0": ev.get("t0", 0.0),
                                "t1": ev.get("t1", 0.0),
                                "text": ev.get("text", ""),
                            })
                        else:
                            _broadcast({
                                "type": "marker",
                                "family": fam,
                                "name": name,
                                "confirmed": bool(ev.get("confirmed")),
                                "t0": ev.get("t0", 0.0),
                                "t1": ev.get("t1", 0.0),
                                "text": ev.get("text", ""),
                            })
                    if session._telemetry:
                        session._telemetry.tick(session._marker_engine.msg_idx)
                        session._telemetry.save(session.id)
                except Exception:
                    pass
        except Exception as e:
            payload = {"type": "error", "error": str(e)}
            _broadcast(payload)
        finally:
            buf.clear()

    while not session._stop_flag:
        try:
            chunk = session._pcm_queue.get(timeout=0.2)
        except queue.Empty:
            # idle-based early flush
            if idle_flush_sec and buf and (time.time() - last_rx) >= idle_flush_sec:
                do_flush()
            continue

        if chunk is FLUSH_SENTINEL:
            do_flush()
            continue

        if not isinstance(chunk, (bytes, bytearray)):
            continue

        if not chunk:
            continue

        buf.extend(chunk)
        last_rx = time.time()
        # periodic soft flush for lower latency (even during continuous speech)
        if soft_target > 0 and len(buf) >= soft_target:
            do_flush()
            continue
        if len(buf) >= target_bytes:
            do_flush()


# ---- API ----
class StartReq(BaseModel):
    lang: Optional[str] = None  # "de" | "en" | None


@app.post("/session/start")
def start_session(req: StartReq):
    sid = str(uuid.uuid4())
    sess = Session(id=sid, lang=req.lang)
    # init diarizer per session based on env/config
    ecapa_dir = os.environ.get("SERAPI_ECAPA_MODEL_DIR")
    ecapa_src = os.environ.get("SERAPI_ECAPA_SOURCE", None)  # e.g., 'speechbrain/spkrec-ecapa-voxceleb'
    # pull defaults from config if envs are not set
    diar_cfg = _CFG.get("diarization", {}) or {}
    os.environ.setdefault("SERAPI_ECAPA_TAU", str(diar_cfg.get("ecapa_tau", 0.25)))
    os.environ.setdefault("SERAPI_DIAR_STICKINESS_DELTA", str(diar_cfg.get("stickiness_delta", 0.03)))
    os.environ.setdefault("SERAPI_DIAR_MIN_TURN_SEC", str(diar_cfg.get("min_turn_sec", 0.6)))
    os.environ.setdefault("SERAPI_DIAR_VAD_ENERGY_DB", str(diar_cfg.get("vad_energy_db", -35)))
    os.environ.setdefault("SERAPI_DIAR_MIN_VOICED_MS", str(diar_cfg.get("min_voiced_ms", 250)))
    os.environ.setdefault("SERAPI_DIAR_MOMENTUM", str(diar_cfg.get("momentum", 0.15)))
    try:
        sess._diarizer = DiarizerFactory.from_env(
            ecapa_dir,
            ecapa_src,
            tau=float(os.environ.get("SERAPI_ECAPA_TAU", "0.25")),
            stickiness_delta=float(os.environ.get("SERAPI_DIAR_STICKINESS_DELTA", "0.03")),
            min_turn_sec=float(os.environ.get("SERAPI_DIAR_MIN_TURN_SEC", "0.6")),
            momentum=float(os.environ.get("SERAPI_DIAR_MOMENTUM", "0.15")),
            vad_energy_db=float(os.environ.get("SERAPI_DIAR_VAD_ENERGY_DB", "-35")),
            min_voiced_ms=int(os.environ.get("SERAPI_DIAR_MIN_VOICED_MS", "250")),
        )
    except Exception:
        sess._diarizer = DiarizerFactory.from_env(None, None)
    # init marker engine + telemetry
    try:
        markers_root = _CFG.get("paths", {}).get("markers_root", "./markers")
        bundle_path = Path("bundles/SerapiCore_1.0.yaml")
        sess._marker_engine = MarkerEngine(MarkerBundle(str(bundle_path), markers_root))
        d = ensure_session_dir(sid)
        sess._telemetry = IntuitionTelemetry(
            bundle_id=sess._marker_engine.bundle.cfg.id,
            data_dir=d,
            scoring_cfg=sess._marker_engine.bundle.cfg.scoring,
        )
    except Exception:
        sess._marker_engine = None
        sess._telemetry = None
    SESSIONS[sid] = sess
    d = ensure_session_dir(sid)
    (d / "meta.json").write_text(json.dumps({"lang": req.lang}, indent=2), encoding="utf-8")
    return {"session_id": sid}


@app.get("/session/{sid}/transcript")
def get_transcript(sid: str):
    sess = SESSIONS.get(sid)
    if not sess:
        return JSONResponse({"error": "session not found"}, status_code=404)
    return {"session_id": sid, "segments": sess.segments}


@app.get("/session/{sid}/export.srt", response_class=PlainTextResponse)
def export_srt(sid: str):
    sess = SESSIONS.get(sid)
    if not sess:
        return PlainTextResponse("session not found", status_code=404)
    srt = segments_to_srt(sess.segments)
    d = ensure_session_dir(sid)
    (d / "export.srt").write_text(srt, encoding="utf-8")
    return srt


@app.get("/session/{sid}/export.txt", response_class=PlainTextResponse)
def export_txt(sid: str, timestamps: bool = False, ts: Optional[str] = None):
    sess = SESSIONS.get(sid)
    if not sess:
        return PlainTextResponse("session not found", status_code=404)
    mode = ts if ts in {"none", "start", "span"} else ("span" if timestamps else "none")
    txt = segments_to_txt(sess.segments, with_timestamps=timestamps, ts_mode=mode)
    d = ensure_session_dir(sid)
    (d / "export.txt").write_text(txt, encoding="utf-8")
    return txt


@app.get("/session/list")
def list_sessions():
    out = []
    for sid, sess in SESSIONS.items():
        out.append({
            "id": sid,
            "lang": sess.lang,
            "segments_count": len(sess.segments),
        })
    return {"sessions": out}


@app.post("/session/{sid}/flush")
def flush_session(sid: str):
    sess = SESSIONS.get(sid)
    if not sess:
        return JSONResponse({"error": "session not found"}, status_code=404)
    # trigger worker flush if running
    try:
        sess._pcm_queue.put(FLUSH_SENTINEL)
        return {"ok": True}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@app.post("/session/{sid}/stop")
def stop_session(sid: str):
    sess = SESSIONS.get(sid)
    if not sess:
        return JSONResponse({"error": "session not found"}, status_code=404)
    sess._stop_flag = True
    return {"ok": True}


@app.websocket("/ws/stream")
async def ws_stream(ws: WebSocket):
    """
    Client verbindet sich mit:
      ws://localhost:8710/ws/stream?session_id=<id>&sr=16000
    und sendet BINARY Frames (PCM16 mono, <i2 little-endian) mit 20–40ms Dauer pro Frame.
    Server streamt 'partial_transcript' Nachrichten als JSON zurück.
    """
    await ws.accept()
    params = ws.query_params
    sid = params.get("session_id")
    try:
        sr = int(params.get("sr", "16000"))
    except Exception:
        sr = 16000
    try:
        window = int(float(params.get("window", str(WINDOW_SECONDS))))
    except Exception:
        window = WINDOW_SECONDS
    try:
        idle_flush = float(params.get("idle_flush", str(IDLE_FLUSH_SECONDS)))
    except Exception:
        idle_flush = IDLE_FLUSH_SECONDS
    try:
        soft_flush = float(params.get("soft_flush", "0"))
    except Exception:
        soft_flush = 0.0

    sess = SESSIONS.get(sid)
    if not sess:
        await ws.close(code=4000)
        return

    # Async queue for worker -> websocket messages
    import asyncio
    out_q: asyncio.Queue = asyncio.Queue(maxsize=32)
    loop = asyncio.get_running_loop()

    # Start worker once per session
    if not sess._worker or not sess._worker.is_alive():
        sess._stop_flag = False

        t = threading.Thread(
            target=_worker_loop, args=(sess, sr, window, loop, out_q, idle_flush, soft_flush), daemon=True
        )
        sess._worker = t
        t.start()

    async def reader_task():
        try:
            while True:
                msg = await ws.receive()
                if (isinstance(msg, dict) and msg.get("bytes")) or (hasattr(msg, "get") and msg.get("bytes")):
                    sess._pcm_queue.put(msg["bytes"])  # handoff to worker
                elif (isinstance(msg, dict) and msg.get("text")) or (hasattr(msg, "get") and msg.get("text")):
                    # optional control commands: 'flush' or JSON {"type":"flush"}
                    txt = msg.get("text")
                    cmd = None
                    if txt:
                        try:
                            maybe = json.loads(txt)
                            cmd = maybe.get("type") if isinstance(maybe, dict) else None
                        except Exception:
                            cmd = txt.strip().lower()
                    if cmd == "flush":
                        sess._pcm_queue.put(FLUSH_SENTINEL)
                        # best-effort ack
                        await ws.send_text(json.dumps({"type": "ack", "cmd": "flush"}))
        except WebSocketDisconnect:
            return
        except Exception:
            return

    async def writer_task():
        try:
            while True:
                payload = await out_q.get()
                await ws.send_text(json.dumps(payload, ensure_ascii=False))
        except WebSocketDisconnect:
            return
        except Exception:
            return

    reader = asyncio.create_task(reader_task())
    writer = asyncio.create_task(writer_task())

    try:
        await asyncio.wait({reader, writer}, return_when=asyncio.FIRST_COMPLETED)
    finally:
        sess._stop_flag = True
        for task in (reader, writer):
            if not task.done():
                task.cancel()
        try:
            await ws.close()
        except Exception:
            pass


@app.get("/healthz")
def health():
    model_dir = str(MODEL_DIR)
    model_dir_exists = os.path.exists(MODEL_DIR)
    model_loaded = whisper_model is not None and not FAKE_ASR
    # diarizer info
    diarizer_mode = "heuristic"
    diarizer_ok = True
    try:
        # peek at any session
        if SESSIONS:
            sess = next(iter(SESSIONS.values()))
            dn = type(sess._diarizer).__name__ if getattr(sess, "_diarizer", None) else "None"
            diarizer_mode = "ecapa" if "Ecapa" in dn else ("heuristic" if "Heuristic" in dn else dn)
    except Exception:
        diarizer_ok = False

    return {
        "ok": True,
        "mode": RUN_MODE,
        "allow_download": ALLOW_DOWNLOAD,
        "model_id": MODEL_ID,
        "model_dir": model_dir,
        "model_dir_exists": model_dir_exists,
        "model_loaded": model_loaded,
        "model_error": MODEL_ERROR,
        "data_root": str(DATA_ROOT),
        "window_seconds": WINDOW_SECONDS,
        "idle_flush_seconds": IDLE_FLUSH_SECONDS,
        "diarizer": {"mode": diarizer_mode, "ok": diarizer_ok},
    }


def _compute_stats(segments: List[Dict[str, Any]]) -> Dict[str, Any]:
    totals: Dict[str, float] = {}
    total_dur = 0.0
    for s in segments:
        dur = max(float(s.get('t1', 0.0)) - float(s.get('t0', 0.0)), 0.0)
        spk = s.get('speaker') or 'UNK'
        totals[spk] = totals.get(spk, 0.0) + dur
        total_dur += dur
    shares = {k: (v / total_dur if total_dur > 0 else 0.0) for k, v in totals.items()}
    return {
        'speakers': [{'speaker': k, 'seconds': round(v, 3), 'share': round(shares[k], 3)} for k, v in sorted(totals.items())],
        'total_seconds': round(total_dur, 3),
    }


@app.get("/session/{sid}/stats")
def get_stats(sid: str):
    sess = SESSIONS.get(sid)
    if not sess:
        return JSONResponse({"error": "session not found"}, status_code=404)
    stats = _compute_stats(sess.segments)
    d = ensure_session_dir(sid)
    (d / 'stats.json').write_text(json.dumps(stats, indent=2), encoding='utf-8')
    return stats


@app.get("/session/{sid}/summary")
def get_summary(sid: str, mode: Optional[str] = None):
    sess = SESSIONS.get(sid)
    if not sess:
        return JSONResponse({"error": "session not found"}, status_code=404)
    summ = Summarizer(mode=mode)
    text = summ.summarize(sess.segments, lang=sess.lang)
    return {"session_id": sid, "mode": summ.mode, "summary": text}


@app.get("/session/{sid}/topics")
def get_topics(sid: str, max_topics: int = 8):
    sess = SESSIONS.get(sid)
    if not sess:
        return JSONResponse({"error": "session not found"}, status_code=404)
    topics = extract_topics(sess.segments, max_topics=max_topics)
    return {"session_id": sid, "topics": topics}


@app.get("/session/{sid}/chapters")
def get_chapters(sid: str, min_segments: int = 6, min_seconds: float = 90.0, drift_threshold: float = 0.55):
    sess = SESSIONS.get(sid)
    if not sess:
        return JSONResponse({"error": "session not found"}, status_code=404)
    chapters = detect_chapters(sess.segments, min_segments=min_segments, min_seconds=min_seconds, drift_threshold=drift_threshold)
    return {"session_id": sid, "chapters": chapters}


@app.get("/session/{sid}/report.md", response_class=PlainTextResponse)
def get_report_md(sid: str, mode: Optional[str] = None, max_topics: int = 8):
    sess = SESSIONS.get(sid)
    if not sess:
        return PlainTextResponse("session not found", status_code=404)
    topics = extract_topics(sess.segments, max_topics=max_topics)
    chapters = detect_chapters(sess.segments)
    summ = Summarizer(mode=mode)
    global_summary = summ.summarize(sess.segments, lang=sess.lang)
    per_ch_summ: list[str] = []
    for ch in chapters:
        i0, i1 = int(ch.get('start_idx', 0)), int(ch.get('end_idx', 0))
        segs = sess.segments[i0:i1]
        per_ch_summ.append(summ.summarize(segs, lang=sess.lang))
    md = render_markdown(sid, sess.segments, topics, chapters, global_summary, per_ch_summ)
    # Append marker interpretation/overview
    try:
        interp_md = _build_marker_interpretation_markdown(sess)
        if interp_md:
            md = md + "\n" + interp_md
    except Exception:
        pass
    d = ensure_session_dir(sid)
    (d / 'report.md').write_text(md, encoding='utf-8')
    return md


def _build_marker_interpretation_markdown(session: 'Session') -> str:
    from engine.marker_engine_core import MarkerBundle, MarkerEngine
    from engine.scoring_engine import IntuitionTelemetry
    # Load bundle fresh to avoid session state interference
    markers_root = _CFG.get("paths", {}).get("markers_root", "./markers")
    bundle_path = Path("bundles/SerapiCore_1.0.yaml")
    me = MarkerEngine(MarkerBundle(str(bundle_path), markers_root))
    tele = IntuitionTelemetry(
        bundle_id=me.bundle.cfg.id,
        data_dir=ensure_session_dir(session.id),
        scoring_cfg=me.bundle.cfg.scoring,
    )
    counts: Dict[str, Dict[str, int]] = {"ATO": {}, "SEM": {}, "CLU": {}}
    confirmed_clu: List[Dict[str, Any]] = []
    # Feed segments sequentially to rebuild events
    all_segs: List[Dict[str, Any]] = []
    for seg in session.segments:
        evs = me.process([seg], all_segs)
        all_segs.append(seg)
        for ev in evs:
            fam = str(ev.get("family", ""))
            name = str(ev.get("name", ""))
            if fam in counts:
                counts[fam][name] = counts[fam].get(name, 0) + 1
            if fam == "CLU" and bool(ev.get("confirmed")):
                confirmed_clu.append(ev)
            if name.startswith("CLU_INTUITION"):
                tele.on_provisional(name, int(ev.get("msg_idx", 0)))
        tele.tick(me.msg_idx)
    states = tele.states
    # Build markdown
    out: List[str] = []
    out.append("## Marker‑Interpretation")
    # High level
    total_ato = sum(counts["ATO"].values())
    total_sem = sum(counts["SEM"].values())
    total_clu = sum(counts["CLU"].values())
    out.append(f"- Gesamt ATO: {total_ato}, SEM: {total_sem}, CLU: {total_clu}")
    # Top examples
    def _top(d: Dict[str, int], k: int = 5) -> List[Tuple[str, int]]:
        return sorted(d.items(), key=lambda x: x[1], reverse=True)[:k]
    top_sem = _top(counts["SEM"]) if counts["SEM"] else []
    top_clu = _top(counts["CLU"]) if counts["CLU"] else []
    if top_sem:
        out.append("- Top SEM: " + ", ".join([f"{n} ({c})" for n, c in top_sem]))
    if top_clu:
        out.append("- Top CLU: " + ", ".join([f"{n} ({c})" for n, c in top_clu]))
    # Confirmed CLU
    if confirmed_clu:
        out.append("- Bestätigte CLU‑Fenster:")
        for ev in confirmed_clu[:10]:
            t0 = float(ev.get("t0", 0.0))
            t1 = float(ev.get("t1", 0.0))
            out.append(f"  - {ev.get('name')} [{t0:.1f}–{t1:.1f}] {ev.get('text','')[:80]}")
    # Intuition states
    if states:
        out.append("- Intuitionen:")
        for name, st in states.items():
            out.append(
                f"  - {name}: provisional={st.get('provisional_hits',0)}, confirmed={st.get('confirmed_hits',0)}, active={bool(st.get('multiplier_active'))}"
            )
    return "\n".join(out)


@app.get("/session/{sid}/report.txt", response_class=PlainTextResponse)
def get_report_txt(sid: str, mode: Optional[str] = None, max_topics: int = 8):
    md = get_report_md(sid, mode=mode, max_topics=max_topics)
    # naive MD -> TXT: strip headings and bullets symbols
    txt = []
    for line in md.splitlines():
        l = line
        l = l.lstrip('#').lstrip().replace("- ", "• ")
        txt.append(l)
    body = "\n".join(txt) + ("\n" if txt else "")
    d = ensure_session_dir(sid)
    (d / 'report.txt').write_text(body, encoding='utf-8')
    return body


@app.get("/session/{sid}/report.docx")
def get_report_docx(sid: str, mode: Optional[str] = None, max_topics: int = 8):
    md = get_report_md(sid, mode=mode, max_topics=max_topics)
    d = ensure_session_dir(sid)
    outp = d / 'report.docx'
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        return JSONResponse({"error": f"python-docx not available: {e}"}, status_code=500)
    doc = Document()
    for line in md.splitlines():
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('- '):
            p = doc.add_paragraph()
            p.add_run('• ').bold = False
            p.add_run(line[2:])
        else:
            doc.add_paragraph(line)
    doc.save(str(outp))
    return JSONResponse({"ok": True, "path": str(outp)})


@app.get("/session/{sid}/report.pdf")
def get_report_pdf(sid: str, mode: Optional[str] = None, max_topics: int = 8):
    md = get_report_md(sid, mode=mode, max_topics=max_topics)
    d = ensure_session_dir(sid)
    outp = d / 'report.pdf'
    try:
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.pdfgen import canvas  # type: ignore
        from reportlab.lib.units import cm  # type: ignore
    except Exception as e:
        return JSONResponse({"error": f"reportlab not available: {e}"}, status_code=500)
    c = canvas.Canvas(str(outp), pagesize=A4)
    width, height = A4
    x = 2 * cm
    y = height - 2 * cm
    for raw in md.splitlines():
        line = raw.replace('\t', '    ')
        if not line:
            y -= 0.5 * cm
        else:
            c.drawString(x, y, line[:120])
            y -= 0.6 * cm
        if y < 2 * cm:
            c.showPage()
            y = height - 2 * cm
    c.save()
    return JSONResponse({"ok": True, "path": str(outp)})


@app.get("/session/{sid}/debug/markers")
def debug_markers(sid: str):
    sess = SESSIONS.get(sid)
    if not sess:
        return JSONResponse({"error": "session not found"}, status_code=404)
    if not sess._marker_engine:
        return {"session_id": sid, "loaded": False}
    me = sess._marker_engine
    bundle = me.bundle
    # summarize examples per marker
    ex_info = {}
    for name in sorted(list(bundle.whitelist())):
        e = bundle.examples.get(name) or {"positive": [], "negative": []}
        ex_info[name] = {"pos": len(e.get("positive", [])), "neg": len(e.get("negative", []))}
    return {
        "session_id": sid,
        "loaded": True,
        "bundle_id": bundle.cfg.id,
        "version": bundle.cfg.version,
        "markers": ex_info,
    }


@app.websocket("/ws/events")
async def ws_events(ws: WebSocket):
    await ws.accept()
    params = ws.query_params
    sid = params.get("session_id")
    sess = SESSIONS.get(sid)
    if not sess:
        await ws.close(code=4000)
        return

    import asyncio
    q: asyncio.Queue = asyncio.Queue(maxsize=64)
    loop = asyncio.get_running_loop()
    sess.add_subscriber(loop, q)

    try:
        while True:
            payload = await q.get()
            await ws.send_text(json.dumps(payload, ensure_ascii=False))
    except WebSocketDisconnect:
        pass
    except Exception:
        pass
    finally:
        sess.remove_subscriber(q)
        try:
            await ws.close()
        except Exception:
            pass


# ---- Report Generation Endpoints ----

@app.post("/session/{sid}/report/generate")
async def generate_session_report(
    sid: str, 
    background_tasks: BackgroundTasks,
    format: str = "markdown",
    include_telemetry: bool = True,
    include_markers: bool = True,
    include_transcript: bool = True,
    include_stats: bool = True
):
    """Generate a report for a session."""
    sess = SESSIONS.get(sid)
    if not sess:
        return JSONResponse({"error": "session not found"}, status_code=404)
    
    try:
        report_format = ReportFormat(format.lower())
    except ValueError:
        return JSONResponse({"error": f"unsupported format: {format}"}, status_code=400)
    
    request = ReportRequest(
        session_id=sid,
        format=report_format,
        include_telemetry=include_telemetry,
        include_markers=include_markers,
        include_transcript=include_transcript,
        include_stats=include_stats
    )
    
    background_tasks.add_task(generate_report_background, report_generator, request)
    
    return {
        "message": "report generation started",
        "session_id": sid,
        "format": format,
        "request_id": f"{sid}_{format}_{int(time.time())}"
    }


@app.get("/session/{sid}/reports")
def list_session_reports(sid: str):
    """List all reports for a session."""
    sess = SESSIONS.get(sid)
    if not sess:
        return JSONResponse({"error": "session not found"}, status_code=404)
    
    reports = report_generator.list_session_reports(sid)
    return {"session_id": sid, "reports": reports}


# ---- Telemetry Endpoints ----

@app.get("/telemetry/health")
def get_telemetry_health():
    """Get system health telemetry."""
    health_data = {
        "timestamp": datetime.utcnow().isoformat(),
        "system": {
            "mode": "mock" if FAKE_ASR else "real",
            "model_loaded": _whisper_pipeline is not None and not FAKE_ASR,
            "active_sessions": len(SESSIONS),
            "data_root": str(DATA_ROOT),
        },
        "sessions": {}
    }
    
    for sid, sess in SESSIONS.items():
        health_data["sessions"][sid] = {
            "segments_count": len(sess.segments),
            "language": sess.lang,
            "has_worker": hasattr(sess, '_worker_thread') and sess._worker_thread is not None
        }
    
    return health_data


@app.get("/telemetry/intuition")
def get_intuition_telemetry():
    """Get CLU_INTUITION marker telemetry."""
    intuition_data = {}
    
    for sid, sess in SESSIONS.items():
        if not sess._marker_engine:
            continue
            
        # Get marker events from the session
        markers = getattr(sess, '_marker_events', [])
        
        for marker in markers:
            name = marker.get('name', '')
            if name.startswith('CLU_INTUITION_'):
                family = name.replace('CLU_INTUITION_', '')
                if family not in intuition_data:
                    intuition_data[family] = {
                        "total_detections": 0,
                        "confirmed": 0,
                        "provisional": 0,
                        "sessions": []
                    }
                
                intuition_data[family]["total_detections"] += 1
                intuition_data[family]["sessions"].append(sid)
                
                # Simple heuristic: consider recent markers as confirmed
                if marker.get('timestamp', 0) > time.time() - 300:  # Last 5 minutes
                    intuition_data[family]["confirmed"] += 1
                else:
                    intuition_data[family]["provisional"] += 1
    
    # Calculate precision estimates
    for family, data in intuition_data.items():
        total = data["total_detections"]
        confirmed = data["confirmed"]
        data["precision"] = (confirmed / total * 100) if total > 0 else 0.0
        data["sessions"] = list(set(data["sessions"]))  # Remove duplicates
    
    return {
        "timestamp": datetime.utcnow().isoformat(),
        "families": intuition_data,
        "summary": {
            "total_families": len(intuition_data),
            "total_detections": sum(d["total_detections"] for d in intuition_data.values()),
            "avg_precision": sum(d["precision"] for d in intuition_data.values()) / len(intuition_data) if intuition_data else 0.0
        }
    }


@app.get("/telemetry/markers/{session_id}")
def get_session_marker_telemetry(session_id: str):
    """Get detailed marker telemetry for a specific session."""
    sess = SESSIONS.get(session_id)
    if not sess:
        return JSONResponse({"error": "session not found"}, status_code=404)
    
    if not sess._marker_engine:
        return JSONResponse({"error": "marker engine not loaded"}, status_code=404)
    
    # Get all marker events for this session
    marker_events = getattr(sess, '_marker_events', [])
    
    # Group by marker type
    marker_stats = {}
    for event in marker_events:
        name = event.get('name', 'unknown')
        if name not in marker_stats:
            marker_stats[name] = {
                "count": 0,
                "first_detection": None,
                "last_detection": None,
                "avg_confidence": 0.0,
                "events": []
            }
        
        marker_stats[name]["count"] += 1
        marker_stats[name]["events"].append(event)
        
        timestamp = event.get('timestamp', 0)
        if marker_stats[name]["first_detection"] is None:
            marker_stats[name]["first_detection"] = timestamp
        marker_stats[name]["last_detection"] = timestamp
        
        # Calculate average confidence if available
        confidence = event.get('confidence', 1.0)
        current_avg = marker_stats[name]["avg_confidence"]
        count = marker_stats[name]["count"]
        marker_stats[name]["avg_confidence"] = (current_avg * (count - 1) + confidence) / count
    
    return {
        "session_id": session_id,
        "timestamp": datetime.utcnow().isoformat(),
        "total_markers": len(marker_events),
        "unique_types": len(marker_stats),
        "markers": marker_stats
    }


# ---- Configuration Validation Endpoint ----

@app.get("/admin/validate-config")
def validate_system_config():
    """Validate system configuration."""
    try:
        from tools.validate_config import ConfigValidator
        
        validator = ConfigValidator()
        valid, results = validator.validate_all()
        
        return {
            "valid": valid,
            "timestamp": datetime.utcnow().isoformat(),
            "results": [
                {
                    "level": r.level.value,
                    "message": r.message,
                    "component": r.component,
                    "details": r.details
                }
                for r in results
            ],
            "summary": {
                "errors": len([r for r in results if r.level.value == "error"]),
                "warnings": len([r for r in results if r.level.value == "warning"]),
                "info": len([r for r in results if r.level.value == "info"])
            }
        }
    except Exception as e:
        logger.handle_error(e, {"endpoint": "validate_config"})
        return JSONResponse(
            {"error": "configuration validation failed", "details": str(e)}, 
            status_code=500
        )


# ---- Static UI ----
try:
    app.mount("/static", StaticFiles(directory="services/static"), name="static")
except Exception:
    pass


@app.get("/")
def ui_root():
    return RedirectResponse(url="/static/index.html")
